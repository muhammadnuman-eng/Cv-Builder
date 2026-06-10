import re
import copy
import os
from docx import Document
from docx.oxml.ns import qn
from generator.services.format_preserver import get_primary_font, get_primary_font_size

SECTION_ORDER = [
    "header", "contact", "summary", "objective",
    "experience", "education", "skills",
    "projects", "certifications", "languages", "awards"
]

HEADING_PATTERNS = {
    "summary":        r"^(summary|objective|about me|profile|professional summary)",
    "experience":     r"^(experience|work experience|employment|career|professional experience)",
    "education":      r"^(education|academic|degree|university|qualification)",
    "skills":         r"^(skills|technical skills|technologies|tech stack|core competencies)",
    "projects":       r"^(projects|key projects|portfolio|personal projects)",
    "certifications": r"^(certifications?|certificates?|courses?|training)",
    "languages":      r"^languages?",
    "awards":         r"^(awards?|achievements?|honors?)",
}


def generate_docx(sections: dict, layout: dict, output_path: str,
                  original_path: str = None, original_format: str = None) -> str:
    if original_path and os.path.exists(original_path) and original_format == "docx":
        return _generate_from_docx_template(original_path, sections, output_path)
    return _generate_fresh(sections, layout, output_path)


# ── In-place modification of original DOCX ──────────────────────────────────

def _generate_from_docx_template(original_path: str, sections: dict, output_path: str) -> str:
    doc = Document(original_path)
    paras = list(doc.paragraphs)

    # Map paragraphs → sections
    section_map = {}   # section_name → {"head": idx, "body": [idx,...]}
    current = None

    for i, para in enumerate(paras):
        text = para.text.strip()
        if not text:
            continue
        detected = _detect_heading(para, text)
        if detected:
            current = detected
            section_map.setdefault(current, {"head": i, "body": []})
        elif current:
            section_map[current]["body"].append(i)

    # Replace body content for each section
    for sec_name, data in section_map.items():
        if sec_name not in sections:
            continue
        new_lines = [l for l in sections[sec_name].split("\n") if l.strip()]
        body_indices = data["body"]
        if not body_indices:
            continue

        # Replace existing paragraph text
        for j, idx in enumerate(body_indices):
            if j < len(new_lines):
                _replace_para_text(paras[idx], new_lines[j])
            else:
                _replace_para_text(paras[idx], "")

        # Add extra lines if tailored content is longer than original
        if len(new_lines) > len(body_indices):
            ref_para = paras[body_indices[-1]]
            prev_elem = ref_para._element
            for extra in new_lines[len(body_indices):]:
                new_elem = copy.deepcopy(ref_para._element)
                _set_element_text(new_elem, extra)
                prev_elem.addnext(new_elem)
                prev_elem = new_elem

    doc.save(output_path)
    return output_path


def _detect_heading(para, text: str) -> str | None:
    style = para.style.name.lower()
    is_heading_style = "heading" in style
    is_all_bold = bool(para.runs) and all(r.bold for r in para.runs if r.text.strip())
    is_upper = text.isupper() and len(text.split()) <= 5

    if not (is_heading_style or is_all_bold or is_upper):
        return None

    normalized = text.lower().strip(":-_ ")
    for section, pattern in HEADING_PATTERNS.items():
        if re.search(pattern, normalized, re.IGNORECASE):
            return section
    return None


def _replace_para_text(para, new_text: str):
    if not para.runs:
        para.add_run(new_text)
        return
    first = para.runs[0]
    first.text = new_text
    for run in para.runs[1:]:
        run.text = ""


def _set_element_text(elem, text: str):
    """Set text in a copied paragraph element: put text in the first run, clear the rest."""
    runs = elem.findall(".//" + qn("w:r"))
    for i, r in enumerate(runs):
        for t in r.findall(qn("w:t")):
            t.text = text if i == 0 else ""


# ── Fresh DOCX (fallback when no original) ──────────────────────────────────

def _generate_fresh(sections: dict, layout: dict, output_path: str) -> str:
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    font_name = get_primary_font(layout)
    font_size = get_primary_font_size(layout)

    doc.styles["Normal"].font.name = font_name
    doc.styles["Normal"].font.size = Pt(font_size)

    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(0.75)

    ordered = [k for k in SECTION_ORDER if k in sections]
    rest    = [k for k in sections if k not in ordered]

    for key in ordered + rest:
        content = sections.get(key, "").strip()
        if not content:
            continue
        if key == "header":
            h = doc.add_heading(content, level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif key == "contact":
            p = doc.add_paragraph(content)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            doc.add_heading(key.upper(), level=2)
            for line in content.split("\n"):
                line = line.strip()
                if line:
                    doc.add_paragraph(f"• {line[1:].strip()}" if line.startswith(("-", "*", "•")) else line)

    doc.save(output_path)
    return output_path
